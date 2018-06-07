import docx

file_name = r'D:\产品\方案沟通\GIS架构设计方案--第一阶段.docx'

doc = docx.Document(file_name)
# data = doc.paragraphs[1].text

print("段落数:"+str(len(doc.paragraphs)))
# 段落数为13，每个回车隔离一段

# 输出每一段的内容
for para in doc.paragraphs:
    print(para.text)


# # 转换成html
from docx2html import convert
import HTMLParser
#
# html_parser = HTMLParser.HTMLParser()
# #使用docx2html模块将docx文件转成html串，随后你想干嘛都行
# html = convert(file_name)
#
# # 这句非常关键，docx2html模块将中文进行了转义，所以要将生成的字符串重新转义
# print(html_parser.enescape(html))